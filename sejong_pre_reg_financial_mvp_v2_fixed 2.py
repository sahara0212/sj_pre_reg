import streamlit as st
import pandas as pd
import io
import altair as alt
from datetime import datetime

# 워드 리포트 생성을 위한 라이브러리 체크 및 임포트
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# -----------------------------
# 0. 세종 Knowledge Base & Impact Data
# -----------------------------

# 규제별 임팩트 점수 (시각화용 데이터: 시급성, 파급력, 복잡도)
REGULATION_IMPACT_DATA = {
    "AI Act": {"urgency": 9, "impact": 9, "complexity": 8, "type": "EU"},
    "GDPR": {"urgency": 7, "impact": 8, "complexity": 7, "type": "EU"},
    "DSA": {"urgency": 6, "impact": 7, "complexity": 6, "type": "EU"},
    "DMA": {"urgency": 5, "impact": 9, "complexity": 9, "type": "EU"},
    "Data Act": {"urgency": 6, "impact": 6, "complexity": 6, "type": "EU"},
    "NIS2": {"urgency": 10, "impact": 10, "complexity": 9, "type": "EU"},
    "CRA": {"urgency": 8, "impact": 8, "complexity": 7, "type": "EU"},
    "PIPA": {"urgency": 9, "impact": 10, "complexity": 8, "type": "KR"},
    "CreditInfo": {"urgency": 8, "impact": 9, "complexity": 9, "type": "KR"},
    "EFTA": {"urgency": 9, "impact": 10, "complexity": 8, "type": "KR"},
    "ISMS-P": {"urgency": 7, "impact": 7, "complexity": 6, "type": "KR"},
    "FCP": {"urgency": 9, "impact": 9, "complexity": 7, "type": "KR"},
    "AML": {"urgency": 10, "impact": 10, "complexity": 7, "type": "KR"},
    "ISMS": {"urgency": 8, "impact": 7, "complexity": 6, "type": "KR"}
}

REGULATIONS_EU = {
    "AI Act": {"kr_term": "EU 인공지능법(AI Act)", "desc": "위험기반 AI 규율. 고위험 AI에 대한 엄격한 통제."},
    "GDPR": {"kr_term": "EU 일반개인정보보호법(GDPR)", "desc": "전 세계 개인정보 컴플라이언스의 표준."},
    "DSA": {"kr_term": "디지털서비스법(DSA)", "desc": "플랫폼 유해 콘텐츠 관리 및 알고리즘 투명성."},
    "DMA": {"kr_term": "디지털시장법(DMA)", "desc": "거대 플랫폼의 독과점 방지 및 공정 경쟁."},
    "Data Act": {"kr_term": "데이터법(Data Act)", "desc": "IoT 데이터의 공정 접근 및 B2B 공유."},
    "DGA": {"kr_term": "데이터거버넌스법(DGA)", "desc": "데이터 중개 및 공익적 활용 거버넌스."},
    "NIS2": {"kr_term": "NIS2 지침 (사이버보안)", "desc": "핵심 인프라의 사이버 회복력 및 경영진 책임 강화."},
    "CRA": {"kr_term": "사이버복원력법(CRA)", "desc": "디지털 제품 전 주기에 걸친 보안요구사항."}
}

REGULATIONS_KR = {
    "PIPA": {"kr_term": "개인정보 보호법", "desc": "전 산업 공통 적용되는 데이터 처리 기본법."},
    "CreditInfo": {"kr_term": "신용정보법", "desc": "마이데이터, 가명정보 결합 등 데이터 금융의 핵심."},
    "EFTA": {"kr_term": "전자금융거래법", "desc": "IT 안전성, 배상책임, 선불업 규율."},
    "ISMS-P": {"kr_term": "ISMS-P 인증", "desc": "정보보호 관리체계 인증."},
    "FCP": {"kr_term": "금융소비자보호법", "desc": "6대 판매원칙 및 징벌적 과징금, 판매 규제."},
    "AML": {"kr_term": "특정금융정보법 (AML)", "desc": "자금세탁방지, 가상자산 신고 등."},
    "ISMS": {"kr_term": "전자금융감독규정", "desc": "망분리, 클라우드 이용 등 기술적 보안 규제."}
}

# -----------------------------
# 1. 시각화 함수 (Enhanced Impact Map)
# -----------------------------
def draw_impact_map(eu_regs, kr_regs):
    data = []
    
    for reg in eu_regs:
        if reg in REGULATION_IMPACT_DATA:
            info = REGULATION_IMPACT_DATA[reg]
            data.append({
                "Regulation": reg,
                "Region": "EU (Global)",
                "Urgency": info['urgency'],
                "Impact": info['impact'],
                "Complexity": info['complexity'] * 80  # 버블 크기 조정
            })
            
    for reg in kr_regs:
        if reg in REGULATION_IMPACT_DATA:
            info = REGULATION_IMPACT_DATA[reg]
            data.append({
                "Regulation": REGULATIONS_KR[reg]['kr_term'],
                "Region": "Korea (Domestic)",
                "Urgency": info['urgency'],
                "Impact": info['impact'],
                "Complexity": info['complexity'] * 80
            })

    if not data:
        return None

    df = pd.DataFrame(data)

    # 기본 차트 설정
    base = alt.Chart(df).encode(
        x=alt.X('Urgency', title='대응 시급성 (Urgency)', scale=alt.Scale(domain=[4, 11])),
        y=alt.Y('Impact', title='비즈니스 파급력 (Biz Impact)', scale=alt.Scale(domain=[4, 11])),
        tooltip=['Regulation', 'Urgency', 'Impact', 'Region']
    )

    # 버블 차트
    points = base.mark_circle(opacity=0.6).encode(
        size=alt.Size('Complexity', title='대응 복잡도', legend=None),
        color=alt.Color('Region', title='규제 권역', scale=alt.Scale(domain=['Korea (Domestic)', 'EU (Global)'], range=['#0055aa', '#cc0000']))
    )

    # 텍스트 레이블 (가독성 향상)
    text = base.mark_text(
        align='left',
        baseline='middle',
        dx=12,
        fontSize=13,
        fontWeight='bold'
    ).encode(
        text='Regulation'
    )

    # 기준선 (사분면 분할)
    rule_x = alt.Chart(pd.DataFrame({'x': [8]})).mark_rule(color='gray', strokeDash=[3,3]).encode(x='x')
    rule_y = alt.Chart(pd.DataFrame({'y': [8]})).mark_rule(color='gray', strokeDash=[3,3]).encode(y='y')

    # 사분면 라벨
    labels = pd.DataFrame([
        {'x': 10.5, 'y': 10.8, 'text': '🚨 최우선 대응 (Critical)'},
        {'x': 10.5, 'y': 4.5, 'text': '⚠️ 단기 집중 (Urgent)'},
        {'x': 4.5, 'y': 10.8, 'text': '🔭 전략적 대비 (Strategic)'},
        {'x': 4.5, 'y': 4.5, 'text': '👀 모니터링 (Monitor)'}
    ])
    
    quadrant_labels = alt.Chart(labels).mark_text(
        align='center', baseline='middle', fontSize=15, fontWeight='bold', color='gray', opacity=0.5
    ).encode(
        x='x', y='y', text='text'
    )

    return (rule_x + rule_y + quadrant_labels + points + text).properties(
        title='Sejong Strategic Regulatory Impact Matrix',
        height=500
    ).interactive()

# -----------------------------
# 2. 마켓 인텔리전스 (규제 예측 로직)
# -----------------------------
def get_market_intelligence():
    intelligence = {
        "eu_trend": """
        **1. EU 규제 동향 (Enforcement Era):** EU는 GDPR, AI Act, NIS2 등 주요 법안의 입법을 완료하고 **'강력한 집행(Enforcement)' 단계**로 진입했습니다. 공급망 보안과 경영진 책임을 묻는 구조로 강화되고 있습니다.
        """,
        "kr_signal": """
        **2. 국내 위기 시그널 (Crisis Signals):** 최근 금융권 클라우드 장애 및 제3자 해킹 사고는 단순 기술 결함을 넘어 **거버넌스 실패**로 인식되고 있습니다. 감독당국은 이에 대해 무관용 원칙을 적용하고 있습니다.
        """,
        "legislation": """
        **3. 입법 동향 (Legislative Moves):** '전자금융거래법 개정안', '개인정보 보호법 2차 개정' 등 **CEO의 책임을 명확히 하고 징벌적 제재**를 포함하는 법안들이 발의되고 있습니다. 핵심은 **'자율 보안'에 상응하는 '사후 책임 강화'**입니다.
        """,
        "prediction": """
        **🚀 세종의 정책 예측 (Forecast):** 향후 규제는 **'사전적 규제'에서 '원칙 중심(Principle-based)의 자율 보안 및 강력한 사후 제재'**로 전환될 것입니다. **'디지털 복원력(Resilience)'** 확보가 감독의 핵심 척도가 될 것입니다.
        """
    }
    return intelligence

# -----------------------------
# 3. 리스크 분석 로직
# -----------------------------
def analyze_service_eu(inputs):
    applicable = []
    notes = []
    risk_notes = []

    if inputs["uses_ai"]:
        level = "고위험(High-risk)" if inputs["ai_function"] in ["신용평가/대출심사", "고용/채용", "의료/진단", "필수서비스 운영"] else "제한적 위험"
        applicable.append("AI Act")
        notes.append(f"**[AI Act]** 귀사의 서비스는 '{level}' 분류 가능성이 높습니다. 투명성 및 인간 감독 체계 정비가 필수입니다.")
        risk_notes.append("**[AI Act 위기]** 시장 진입 차단 및 최대 전 세계 매출 7% 과징금. 알고리즘 차별 이슈로 인한 집단 소송 리스크.")

    if inputs["is_critical_service"] or inputs["inst_type"] in ["은행", "보험", "증권/투자"]:
        applicable.append("NIS2")
        notes.append("**[NIS2 / DORA]** 금융권은 ICT 리스크 관리 핵심 대상입니다. 경영진(C-Level)의 법적 책임을 강화하는 거버넌스가 시급합니다.")
        risk_notes.append("**[NIS2 위기]** 경영진 해임 권고 및 형사 책임 추궁 가능. EU 내 금융 라이선스 정지 등 치명적 제재.")

    if inputs["processes_personal_data"]:
        applicable.append("GDPR")
        risk_notes.append("**[GDPR 위기]** 최대 매출 4% 과징금 및 글로벌 파트너십 단절 위험.")

    if inputs["is_online_platform"]:
        applicable.append("DSA")
        risk_notes.append("**[DSA 위기]** 불법 콘텐츠 관리 실패 시 반복적 위반으로 간주, 서비스 차단(Shut-down) 조치 가능.")

    return list(set(applicable)), notes, risk_notes

def analyze_service_kr(inputs):
    applicable = []
    notes = []
    risk_notes = []

    if inputs["handles_credit_info"] or inputs["mydata_business"]:
        applicable.append("CreditInfo")
        notes.append("**[신용정보법]** 마이데이터/가명정보 결합 시, 금융보안원 기준 상회 보안 조치 및 전송요구권 대응 체계 필수.")
        risk_notes.append("**[신용정보법 위기]** 데이터 오남용 시 '징벌적 손해배상(3배)', 본허가 취소 및 신규 사업 영구 제한.")

    if inputs["offers_electronic_financial_services"] or inputs["uses_cloud_outsourcing"]:
        applicable.append("EFTA")
        applicable.append("ISMS")
        notes.append("**[전자금융감독규정]** 망분리 완화/SaaS 도입은 혁신이자 리스크입니다. '제3자 리스크' 관리 실패는 중대 과실입니다.")
        risk_notes.append("**[전금법 위기]** 클라우드 장애로 인한 서비스 중단 시, 관리 소홀 입증되면 기관 중징계 및 CEO 문책 경고.")

    if inputs["offers_retail_products"]:
        applicable.append("FCP")
        notes.append("**[금융소비자보호법]** 디지털 채널 UX/UI가 '다크 패턴'으로 오인되지 않도록 설계 단계부터 법률 검토가 필요합니다.")
        risk_notes.append("**[금소법 위기]** 불완전판매 간주 시 판매액 최대 50% 징벌적 과징금 및 대규모 환불 사태.")

    if inputs["aml_high_risk"]:
        applicable.append("AML")
        risk_notes.append("**[특금법 위기]** 자금세탁 방지 미흡은 해외 당국의 Secondary Boycott 대상이 되며, 이는 글로벌 금융망 퇴출을 의미합니다.")

    return list(set(applicable)), notes, risk_notes

# -----------------------------
# 4. 5W1H 액션플랜
# -----------------------------
def build_expert_advisory(eu_regs, kr_regs, inputs):
    actions_fin = []
    actions_reg = []

    if "AI Act" in eu_regs or inputs["uses_ai"]:
        actions_fin.append({
            "area": "AI Governance",
            "title": "AI 리스크 등급 분류 및 설명가능성(XAI) 입증",
            "target": "CRO / CISO",
            "what": "운영 중인 AI 모델의 '리스크 등급' 매핑 및 설명가능성 보고서 작성.",
            "why": "규제 확정 후 대응은 늦습니다. '블랙박스 AI'는 법적 분쟁 시 방어권을 무력화시킵니다.",
            "how": "세종의 'AI 법률/기술 복합 진단'을 통해 모델 명세서를 법적 문서화하고 이사회에 정기 보고하십시오."
        })

    if "EFTA" in kr_regs or "ISMS" in kr_regs:
        actions_fin.append({
            "area": "Digital Resilience",
            "title": "망분리 규제 샌드박스 및 제로 트러스트 전환",
            "target": "CIO / CISO",
            "what": "획일적 망분리에서 '데이터 중요도' 기반 논리적 망분리/제로 트러스트로 전환 로드맵 수립.",
            "why": "물리적 망분리는 AI/클라우드 도입의 걸림돌입니다. 혁신을 위해선 규제 예외를 선제적으로 확보해야 합니다.",
            "how": "세종이 감독당국 소통을 지원하여 '혁신금융서비스' 지정을 돕고, 책임 공유 모델 계약을 재구성해 드립니다."
        })

    if "FCP" in kr_regs:
        actions_fin.append({
            "area": "Consumer Protection",
            "title": "디지털 채널 'Product Governance' 재설계",
            "target": "CCO",
            "what": "앱/웹 상품 가입 동선 내 '다크 패턴' 요소 전수 진단 및 제거.",
            "why": "알고리즘 추천은 '설명의무 위반'의 강력한 증거가 될 수 있습니다.",
            "how": "상품 기획 단계부터 법률 검토를 강제하는 절차를 도입하여 제재 리스크를 차단하십시오."
        })

    # 정책 제언
    actions_reg.append({
        "area": "Regulatory Sandbox",
        "title": "금융권 AI 도입 활성화를 위한 '규제 프리존' 확대",
        "what": "국내 현실에 맞는 '유연한 위험 기반 규율(Risk-based approach)' 가이드라인 제정.",
        "why": "과도한 사전 규제는 경쟁력을 저하시킵니다. 사후 책임은 강화하되 진입 장벽은 낮추는 전환이 필요합니다."
    })
    actions_reg.append({
        "area": "Cloud Security",
        "title": "물리적 망분리 규제의 단계적 폐지 및 원칙 중심 전환",
        "what": "자체 보안 역량 입증 시 물리적 망분리 의무를 면제하는 '자율 보안 인증제' 도입.",
        "why": "글로벌 트렌드(DORA)와 같이 '회복력(Resilience)' 중심 감독으로 전환해야 합니다."
    })

    return actions_fin, actions_reg

# -----------------------------
# 5. 텍스트 & 워드 리포트 생성 함수
# -----------------------------
def generate_report_text(inputs, kr_regs, kr_notes, kr_risks, eu_regs, eu_notes, actions_fin, actions_reg):
    intel = get_market_intelligence()
    lines = []
    lines.append("[법무법인 세종 Regulatory Foresight Report]")
    lines.append(f"발행일: {datetime.now().strftime('%Y-%m-%d')}")
    lines.append("작성자: 법무법인 세종 금융규제그룹 (Digital Finance & Regulatory Group)")
    lines.append("수신: 귀사 경영진 및 이사회")
    lines.append("=" * 70)
    
    lines.append("\n1. Executive Summary: 규제 환경 전망 및 전략적 방향")
    lines.append(f"귀 기관({inputs['inst_type']})은 '{inputs['biz_type']}' 업무와 관련하여 복합적인 '이중 규제' 리스크에 노출되어 있습니다.")
    lines.append("세종은 글로벌 트렌드와 국내 정책 시그널을 분석하여, 귀사가 '규제 준수'를 넘어 '시장 선도'를 할 수 있도록 지원합니다.")

    lines.append("\n2. Regulatory Market Intelligence (규제 및 정책 환경 분석)")
    lines.append(intel['eu_trend'].strip().replace("**", ""))
    lines.append(intel['kr_signal'].strip().replace("**", ""))
    lines.append(intel['legislation'].strip().replace("**", ""))
    lines.append(intel['prediction'].strip().replace("**", ""))

    lines.append("\n3. [Critical Warning] 준비하지 않을 경우 직면할 위기")
    lines.append("선제적 대응 실패 시 예상되는 법적/비즈니스 위기입니다:")
    for risk in kr_risks:
        lines.append(f"- {risk.replace('**', '')}")
    for risk in risk_notes_eu(inputs)[2]:
        lines.append(f"- {risk.replace('**', '')}")

    lines.append("\n4. Sejong's Strategic Action Plan (5W1H)")
    lines.append("[금융회사 경영진을 위한 제언]")
    for action in actions_fin:
        lines.append(f"▶ {action['title']}")
        lines.append(f"   (Why) {action['why']}")
        lines.append(f"   (What) {action['what']}")
        lines.append(f"   (How) {action['how']}")
        lines.append("")

    lines.append("[감독당국 대상 정책 제언 (Policy Advocacy)]")
    for action in actions_reg:
        lines.append(f"▶ {action['title']}")
        lines.append(f"   - {action['what']}")

    lines.append("\n5. Conclusion")
    lines.append("법무법인 세종은 금융위·금감원 출신 전문가들의 통찰력으로 귀사의 미래 리스크를 예방하고 위기관리를 선도하겠습니다.")
    
    return "\n".join(lines)

def create_word_report(inputs, kr_risks, eu_risks, intel, actions_fin, actions_reg):
    if not DOCX_AVAILABLE:
        return None

    doc = Document()
    
    # 1. 표지 및 헤더
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = "법무법인 세종 | SHIN & KIM LLC"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    title = doc.add_heading('Regulatory Foresight Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"발행일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    doc.add_paragraph("작성자: 법무법인 세종 금융규제그룹 (Digital Finance & Regulatory Group)")
    doc.add_paragraph(f"수신: {inputs['inst_type']} 경영진 및 이사회 귀중")
    doc.add_paragraph("-" * 80)

    # 2. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        f"귀 기관({inputs['inst_type']})은 '{inputs['biz_type']}' 업무와 관련하여 디지털 전환 가속화에 따른 "
        "복합적인 '이중 규제(Dual-Regulation)' 리스크에 노출되어 있습니다. "
        "본 보고서는 글로벌 트렌드와 국내 정책 시그널을 종합 분석하여, 귀사가 '규제 준수'를 넘어 "
        "'시장 선도'를 할 수 있는 선제적 대응 전략을 제시합니다."
    )

    # 3. Market Intelligence
    doc.add_heading('2. 규제 및 정책 환경 심층 분석', level=1)
    
    doc.add_heading('2.1 EU 및 글로벌 규제 동향', level=2)
    doc.add_paragraph(intel['eu_trend'].strip().replace("**", ""))
    
    doc.add_heading('2.2 국내 보안 사고 및 위기 시그널', level=2)
    doc.add_paragraph(intel['kr_signal'].strip().replace("**", ""))
    
    doc.add_heading('2.3 국회 입법 및 정책 발의 동향', level=2)
    doc.add_paragraph(intel['legislation'].strip().replace("**", ""))
    
    # 예측 박스 (강조)
    forecast_para = doc.add_paragraph()
    forecast_run = forecast_para.add_run(intel['prediction'].strip().replace("**", ""))
    forecast_run.bold = True
    forecast_run.font.color.rgb = RGBColor(0, 51, 102) # 남색 강조

    # 4. Critical Warning
    doc.add_heading('3. [Critical Warning] 미대응 시 직면할 위기', level=1)
    doc.add_paragraph("선제적 대응 실패 시 예상되는 법적 제재 및 비즈니스 위기입니다:", style='Intense Quote')
    
    risk_list = kr_risks + eu_risks
    for risk in risk_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(risk.replace("**", ""))
        run.font.color.rgb = RGBColor(192, 0, 0) # 붉은색 경고

    # 5. Action Plan
    doc.add_heading('4. 세종 Strategic Action Plan', level=1)
    
    doc.add_heading('4.1 금융회사 경영진을 위한 제언 (Action Items)', level=2)
    
    # 표 생성
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '대응 영역'
    hdr_cells[1].text = '실행 과제 (What)'
    hdr_cells[2].text = '세종 솔루션 (How)'
    
    for action in actions_fin:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{action['area']}\n({action['target']})"
        row_cells[1].text = f"{action['title']}\n\n[Why] {action['why']}"
        row_cells[2].text = action['how']

    doc.add_heading('4.2 감독당국 대상 정책 제언 (Policy Advocacy)', level=2)
    doc.add_paragraph("법무법인 세종은 귀사의 입장을 대변하여 합리적인 규제 환경 조성을 위해 다음과 같이 제언합니다.")
    for action in actions_reg:
        doc.add_paragraph(f"▶ {action['title']}: {action['what']}", style='List Number')

    # 6. Conclusion
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "미래의 사고는 '수습'하는 것이 아니라 '예방'하는 것입니다. "
        "법무법인 세종은 금융위·금감원 출신 전문가들의 통찰력으로 귀사의 미래 리스크를 예방하고 "
        "디지털 금융 시장을 선도할 수 있도록 가장 앞에서 돕겠습니다."
    )

    # 메모리 버퍼에 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def risk_notes_eu(inputs):
    return analyze_service_eu(inputs)

# -----------------------------
# 6. Streamlit UI 메인
# -----------------------------
def main():
    st.set_page_config(
        page_title="법무법인 세종 Regulatory Foresight",
        page_icon="⚖️",
        layout="wide"
    )

    # 상단 헤더
    st.title("법무법인 세종 Regulatory Foresight")
    st.markdown("### 글로벌 규제 예측 및 선제적 위기관리 솔루션")
    st.info(
        """
        **법무법인 세종 금융규제그룹**은 사고 후 대응(Post-Event Response)에 머무르지 않습니다.  
        **글로벌 규제 트렌드와 국내 정책/사고 시그널을 조기에 심층 분석**하여,  
        규제 정책 방향을 **미리 예측(Forecast)**하고 감독기관과 금융기관에 **선제적인 제언(Pre-emptive Advisory)**을 제공함으로써,  
        귀사의 미래 사고 예방과 위기 관리를 **선도**합니다.
        """
    )
    st.divider()

    # Sidebar 입력
    with st.sidebar:
        st.header("🏢 기관 프로파일링")
        inst_type = st.selectbox("기관 유형", ["은행", "증권/투자", "보험", "카드/캐피탈", "핀테크/BigTech", "가상자산사업자"], index=0)
        biz_type = st.selectbox("핵심 사업 영역", 
            ["여신/수신", "지급결제/송금", "자산관리/Robo-Advisor", "보험 인수/클레임", "마이데이터/플랫폼", "가상자산 거래/수탁", "AML/FDS"], index=0)
        
        st.header("🌐 디지털 혁신 & 리스크 요인")
        eu_business = st.checkbox("EU/글로벌 시장 진출 (예정 포함)")
        uses_ai = st.checkbox("AI/알고리즘 도입 (신용평가, 챗봇 등)")
        ai_function = st.selectbox("AI 활용 목적", ["해당없음", "신용평가/심사", "고객상담(챗봇)", "마케팅 추천", "이상거래탐지", "채용/인사"], disabled=not uses_ai)
        
        processes_personal_data = st.checkbox("대량의 개인신용정보 처리")
        uses_cloud_outsourcing = st.checkbox("클라우드(SaaS) 및 외부 위탁 활용")
        is_online_platform = st.checkbox("플랫폼 비즈니스 (중개/매칭)")
        is_critical_service = st.checkbox("핵심 금융 인프라/중요 업무 해당")
        
        aml_high_risk = st.checkbox("자금세탁 고위험(가상자산, 해외송금 등) 취급")
        offers_retail_products = st.checkbox("금융소비자 대상 상품 판매")

        analyze_btn = st.button("규제 예측 및 전략 분석 시작", type="primary")

    if analyze_btn:
        # 입력 데이터 패키징
        inputs = {
            "inst_type": inst_type, "biz_type": biz_type, "eu_business": eu_business,
            "uses_ai": uses_ai, "ai_function": ai_function,
            "processes_personal_data": processes_personal_data,
            "uses_cloud_outsourcing": uses_cloud_outsourcing,
            "is_online_platform": is_online_platform,
            "is_critical_service": is_critical_service,
            "aml_high_risk": aml_high_risk,
            "offers_retail_products": offers_retail_products,
            "uses_sensitive_data": True, 
            "uses_biometric": False, 
            "offers_electronic_financial_services": True,
            "is_very_large_platform": False,
            "provides_software_product": False,
            "shares_iot_data": False,
            "handles_credit_info": True,
            "mydata_business": (biz_type == "마이데이터/플랫폼")
        }

        # 분석 실행
        eu_regs, eu_notes, eu_risks = analyze_service_eu(inputs)
        kr_regs, kr_notes, kr_risks = analyze_service_kr(inputs)
        actions_fin, actions_reg = build_expert_advisory(eu_regs, kr_regs, inputs)
        intel = get_market_intelligence()

        # Tab 구성
        tab1, tab2, tab3, tab4 = st.tabs(["📊 규제 임팩트 맵 & 트렌드 분석", "⚠️ Critical Risk Warning", "💡 세종 Foresight & Action", "📑 Executive Report"])

        with tab1:
            st.subheader("Regulatory Impact & Market Intelligence")
            
            # 1. 트렌드 분석 (텍스트)
            with st.expander("🌍 글로벌/국내 규제 환경 및 정책 시그널 분석 (Click to Expand)", expanded=True):
                col_intel1, col_intel2 = st.columns(2)
                with col_intel1:
                    st.markdown(intel["eu_trend"])
                    st.markdown(intel["legislation"])
                with col_intel2:
                    st.markdown(intel["kr_signal"])
                    st.success(intel["prediction"])
            
            st.divider()

            # 2. 임팩트 맵 (차트)
            st.markdown("#### 🎯 Regulatory Impact Map")
            st.markdown("**X축: 대응 시급성(Urgency) | Y축: 비즈니스 파급력(Impact) | 버블 크기: 대응 복잡도(Complexity)**")
            
            chart = draw_impact_map(eu_regs, kr_regs)
            if chart:
                st.altair_chart(chart, use_container_width=True)
            else:
                st.write("해당되는 주요 규제 이슈가 식별되지 않았습니다.")

        with tab2:
            st.subheader("🛑 [Critical Warning] 준비하지 않을 경우 직면할 위기")
            st.markdown("단순한 과태료를 넘어, **비즈니스의 존폐를 위협할 수 있는 핵심 리스크**를 진단했습니다.")
            
            st.markdown(
                """
                <style>
                .risk-box {
                    background-color: #fff3cd;
                    border-left: 5px solid #ffc107;
                    padding: 15px;
                    border-radius: 5px;
                    margin-bottom: 10px;
                }
                .risk-header {
                    color: #856404;
                    font-weight: bold;
                    font-size: 1.1em;
                }
                </style>
                """, unsafe_allow_html=True
            )

            if kr_risks:
                st.markdown("#### 🇰🇷 Domestic Risks")
                for risk in kr_risks:
                    st.markdown(f'<div class="risk-box"><span class="risk-header">⚠ 경고:</span> {risk.replace("**", "")}</div>', unsafe_allow_html=True)
            
            if eu_risks:
                st.markdown("#### 🇪🇺 Global Risks")
                for risk in eu_risks:
                    st.markdown(f'<div class="risk-box"><span class="risk-header">⚠ 경고:</span> {risk.replace("**", "")}</div>', unsafe_allow_html=True)

        with tab3:
            st.subheader("💡 세종's Strategic Foresight")
            st.markdown("금융위·금감원 출신 전문위원들이 제안하는 **선제적 대응 전략(Action Plan)**입니다.")
            
            st.markdown("#### 1. 금융회사 경영진을 위한 Strategic Action")
            for action in actions_fin:
                with st.expander(f"📌 {action['title']}", expanded=True):
                    st.markdown(f"**WHY (시그널 분석):** {action['why']}")
                    st.markdown(f"**WHAT (대응 과제):** {action['what']}")
                    st.success(f"🚀 **Sejong's Solution:** {action['how']}")

            st.markdown("---")
            st.markdown("#### 2. 감독당국 대상 정책 제언 (Policy Advocacy)")
            st.info("※ 세종은 귀사의 입장을 대변하여 합리적인 규제 환경 조성을 위한 정책 제언을 주도합니다.")
            for action in actions_reg:
                st.markdown(f"**🗣️ {action['title']}**")
                st.caption(f"제언 내용: {action['what']}")

        with tab4:
            st.subheader("📑 Executive Report 생성")
            st.markdown("이사회 및 경영진 보고를 위한 **전문가 리포트**를 생성합니다.")
            
            report_text = generate_report_text(inputs, kr_regs, kr_notes, kr_risks, eu_regs, eu_notes, actions_fin, actions_reg)
            
            st.text_area("리포트 미리보기", report_text, height=500)
            
            st.download_button(
                label="📄 리포트 다운로드 (TXT)",
                data=report_text,
                file_name=f"Sejong_Foresight_Report_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain"
            )
            
            if DOCX_AVAILABLE:
                docx_file = create_word_report(inputs, kr_risks, eu_risks, intel, actions_fin, actions_reg)
                st.download_button(
                    label="📄 공식 리포트 다운로드 (DOCX)",
                    data=docx_file,
                    file_name=f"Sejong_Foresight_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("⚠️ 서버에 'python-docx' 라이브러리가 설치되어 있지 않아 워드 파일을 생성할 수 없습니다.")

    else:
        st.info("👈 좌측 사이드바에서 기관 정보를 입력하고 분석 버튼을 눌러주세요.")
        st.markdown(
            """
            #### 🏆 Why 법무법인 세종 Regulatory Foresight?
            - **Supervisory Insight:** 금융위/금감원 출신 전문가들의 정책 시그널 조기 포착
            - **Global Standard:** EU/미국 로펌과의 협업을 통한 글로벌 규제 트렌드 실시간 반영
            - **Tech-Legal Expertise:** IT/보안 실무 경험을 갖춘 전문위원들의 실질적인 기술적 해법 제시
            """
        )

if __name__ == "__main__":
    main()
